import requests
import pandas as pd
import re
import os
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
import pickle
import io
from google.colab import drive
from google.auth.transport.requests import Request
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.oauth2.credentials import Credentials
from docx import Document
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from datetime import datetime, timedelta

# Mount Google Drive to access files
drive.mount('/content/drive')

# Install required packages
!pip install python-docx google-api-python-client google-auth-httplib2 google-auth-oauthlib beautifulsoup4 nltk

# Download NLTK data
nltk.download('punkt')
nltk.download('stopwords')

# 1. Job Retrieval Component
class JobRetriever:
    def __init__(self):
        self.jobs_df = pd.DataFrame(columns=['Title', 'Company', 'Location', 'Description', 'URL', 'Date_Posted', 'Keywords_Match'])
    
    def search_indeed(self, keywords, location, pages=5):
        """
        Scrape job listings from Indeed based on keywords and location
        """
        print(f"Searching Indeed for {keywords} in {location}...")
        base_url = "https://www.indeed.com/jobs"
        
        all_jobs = []
        for page in range(pages):
            params = {
                'q': keywords,
                'l': location,
                'start': page * 10
            }
            
            # Add user agent to avoid being blocked
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            try:
                response = requests.get(base_url, params=params, headers=headers)
                soup = BeautifulSoup(response.text, 'html.parser')
                
                job_cards = soup.find_all('div', class_='jobsearch-SerpJobCard')
                
                for card in job_cards:
                    job_title_elem = card.find('a', class_='jobtitle')
                    company_elem = card.find('span', class_='company')
                    location_elem = card.find('div', class_='recJobLoc')
                    description_elem = card.find('div', class_='summary')
                    
                    if job_title_elem and company_elem:
                        job_title = job_title_elem.text.strip()
                        company = company_elem.text.strip()
                        location = location_elem['data-rc-loc'] if location_elem else "N/A"
                        description = description_elem.text.strip() if description_elem else "N/A"
                        url = "https://www.indeed.com" + job_title_elem['href']
                        
                        # Calculate keyword match score
                        keywords_list = keywords.lower().split()
                        match_score = sum(1 for keyword in keywords_list if keyword.lower() in (job_title.lower() + " " + description.lower()))
                        
                        job_data = {
                            'Title': job_title,
                            'Company': company,
                            'Location': location,
                            'Description': description,
                            'URL': url,
                            'Date_Posted': datetime.now().strftime("%Y-%m-%d"),
                            'Keywords_Match': match_score
                        }
                        
                        all_jobs.append(job_data)
            
            except Exception as e:
                print(f"Error scraping Indeed page {page}: {str(e)}")
        
        # Update the jobs dataframe
        new_jobs_df = pd.DataFrame(all_jobs)
        self.jobs_df = pd.concat([self.jobs_df, new_jobs_df], ignore_index=True)
    
    def search_linkedin(self, keywords, location, pages=2):
        """
        Scrape job listings from LinkedIn based on keywords and location
        Note: LinkedIn has more restrictions, so this may require additional auth
        """
        print(f"Searching LinkedIn for {keywords} in {location}...")
        # This is a simplified version; LinkedIn may require authentication
        base_url = "https://www.linkedin.com/jobs/search"
        
        all_jobs = []
        for page in range(pages):
            params = {
                'keywords': keywords,
                'location': location,
                'start': page * 25
            }
            
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            try:
                response = requests.get(base_url, params=params, headers=headers)
                soup = BeautifulSoup(response.text, 'html.parser')
                
                job_cards = soup.find_all('div', class_='base-card')
                
                for card in job_cards:
                    job_title_elem = card.find('h3', class_='base-search-card__title')
                    company_elem = card.find('h4', class_='base-search-card__subtitle')
                    location_elem = card.find('span', class_='job-search-card__location')
                    link_elem = card.find('a', class_='base-card__full-link')
                    
                    if job_title_elem and company_elem:
                        job_title = job_title_elem.text.strip()
                        company = company_elem.text.strip()
                        location = location_elem.text.strip() if location_elem else "N/A"
                        url = link_elem['href'] if link_elem else "N/A"
                        
                        # Get full description by visiting the job URL
                        description = "N/A"
                        if url != "N/A":
                            try:
                                job_response = requests.get(url, headers=headers)
                                job_soup = BeautifulSoup(job_response.text, 'html.parser')
                                description_elem = job_soup.find('div', class_='description__text')
                                if description_elem:
                                    description = description_elem.text.strip()
                            except:
                                pass
                        
                        # Calculate keyword match score
                        keywords_list = keywords.lower().split()
                        match_score = sum(1 for keyword in keywords_list if keyword.lower() in (job_title.lower() + " " + description.lower()))
                        
                        job_data = {
                            'Title': job_title,
                            'Company': company,
                            'Location': location,
                            'Description': description,
                            'URL': url,
                            'Date_Posted': datetime.now().strftime("%Y-%m-%d"),
                            'Keywords_Match': match_score
                        }
                        
                        all_jobs.append(job_data)
            
            except Exception as e:
                print(f"Error scraping LinkedIn page {page}: {str(e)}")
        
        # Update the jobs dataframe
        new_jobs_df = pd.DataFrame(all_jobs)
        self.jobs_df = pd.concat([self.jobs_df, new_jobs_df], ignore_index=True)
    
    def filter_jobs(self, min_keywords_match=2):
        """Filter jobs based on keyword match score"""
        return self.jobs_df[self.jobs_df['Keywords_Match'] >= min_keywords_match].sort_values('Keywords_Match', ascending=False)
    
    def save_jobs(self, filename='job_listings.csv'):
        """Save job listings to CSV"""
        path = '/content/drive/My Drive/' + filename
        self.jobs_df.to_csv(path, index=False)
        print(f"Saved {len(self.jobs_df)} jobs to {path}")
        return path

# 2. Resume and Cover Letter Customizer
class DocumentCustomizer:
    def __init__(self, resume_path, cover_letter_path):
        """
        Initialize with paths to resume and cover letter templates
        """
        self.resume_path = resume_path
        self.cover_letter_path = cover_letter_path
        self.resume_doc = None
        self.cover_letter_doc = None
        self.load_documents()
    
    def load_documents(self):
        """Load the Word documents"""
        try:
            self.resume_doc = Document(self.resume_path)
            self.cover_letter_doc = Document(self.cover_letter_path)
            print("Documents loaded successfully")
        except Exception as e:
            print(f"Error loading documents: {str(e)}")
    
    def extract_job_keywords(self, job_description):
        """Extract important keywords from job description"""
        # Tokenize and preprocess
        tokens = word_tokenize(job_description.lower())
        stop_words = set(stopwords.words('english'))
        filtered_tokens = [word for word in tokens if word.isalnum() and word not in stop_words]
        
        # Use TF-IDF to find important terms
        vectorizer = TfidfVectorizer(max_features=20)
        tfidf_matrix = vectorizer.fit_transform([' '.join(filtered_tokens)])
        feature_names = vectorizer.get_feature_names_out()
        
        # Get top keywords
        dense = tfidf_matrix.todense()
        words_importance = dense[0].tolist()[0]
        word_scores = [(word, words_importance[i]) for i, word in enumerate(feature_names)]
        word_scores.sort(key=lambda x: x[1], reverse=True)
        
        return [word for word, score in word_scores[:10]]
    
    def customize_resume(self, job_title, company_name, job_description):
        """Customize resume based on job details"""
        if not self.resume_doc:
            print("Resume document not loaded")
            return None
        
        # Create a copy of the resume
        custom_resume = Document()
        
        # Extract keywords from job description
        keywords = self.extract_job_keywords(job_description)
        print(f"Keywords extracted: {keywords}")
        
        # Customize the resume sections
        for para in self.resume_doc.paragraphs:
            # Add job-specific objective
            if "[OBJECTIVE]" in para.text:
                custom_text = para.text.replace("[OBJECTIVE]", 
                    f"Experienced professional seeking the {job_title} position at {company_name}, bringing expertise in " + 
                    ", ".join(keywords[:3]) + ".")
                custom_resume.add_paragraph(custom_text, para.style)
            else:
                # Highlight keywords in the resume
                text = para.text
                for keyword in keywords:
                    if keyword.lower() in text.lower() and len(keyword) > 3:  # Only highlight meaningful words
                        pattern = re.compile(re.escape(keyword), re.IGNORECASE)
                        text = pattern.sub(f"**{keyword}**", text)
                
                custom_resume.add_paragraph(text, para.style)
        
        # Save the customized resume
        filename = f"Custom_Resume_{company_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        save_path = f"/content/drive/My Drive/{filename}"
        custom_resume.save(save_path)
        print(f"Customized resume saved to {save_path}")
        
        return save_path
    
    def customize_cover_letter(self, job_title, company_name, job_description):
        """Customize cover letter based on job details"""
        if not self.cover_letter_doc:
            print("Cover letter document not loaded")
            return None
        
        # Create a copy of the cover letter
        custom_cover_letter = Document()
        
        # Extract keywords from job description
        keywords = self.extract_job_keywords(job_description)
        
        # Current date
        current_date = datetime.now().strftime("%B %d, %Y")
        
        # Customize the cover letter
        for para in self.cover_letter_doc.paragraphs:
            text = para.text
            
            # Replace placeholders
            replacements = {
                "[DATE]": current_date,
                "[COMPANY_NAME]": company_name,
                "[JOB_TITLE]": job_title,
                "[KEYWORD1]": keywords[0] if keywords else "",
                "[KEYWORD2]": keywords[1] if len(keywords) > 1 else "",
                "[KEYWORD3]": keywords[2] if len(keywords) > 2 else ""
            }
            
            for placeholder, value in replacements.items():
                text = text.replace(placeholder, value)
            
            custom_cover_letter.add_paragraph(text, para.style)
        
        # Save the customized cover letter
        filename = f"Cover_Letter_{company_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        save_path = f"/content/drive/My Drive/{filename}"
        custom_cover_letter.save(save_path)
        print(f"Customized cover letter saved to {save_path}")
        
        return save_path

# 3. Google Calendar Integration for Follow-up Reminders
class CalendarManager:
    def __init__(self):
        """Initialize Google Calendar API"""
        self.creds = None
        self.SCOPES = ['https://www.googleapis.com/auth/calendar']
        self.credentials_path = '/content/drive/My Drive/credentials.json'
        self.token_path = '/content/drive/My Drive/token.pickle'
        self.service = None
        self.authenticate()
    
    def authenticate(self):
        """Authenticate with Google Calendar API"""
        try:
            if os.path.exists(self.token_path):
                with open(self.token_path, 'rb') as token:
                    self.creds = pickle.load(token)
            
            # If credentials don't exist or are invalid
            if not self.creds or not self.creds.valid:
                if self.creds and self.creds.expired and self.creds.refresh_token:
                    self.creds.refresh(Request())
                else:
                    flow = InstalledAppFlow.from_client_secrets_file(
                        self.credentials_path, self.SCOPES)
                    self.creds = flow.run_local_server(port=0)
                
                # Save the credentials for the next run
                with open(self.token_path, 'wb') as token:
                    pickle.dump(self.creds, token)
            
            self.service = build('calendar', 'v3', credentials=self.creds)
            print("Successfully authenticated with Google Calendar")
        
        except Exception as e:
            print(f"Error authenticating with Google Calendar: {str(e)}")
    
    def set_follow_up_reminder(self, job_title, company_name, application_date=None, follow_up_days=5):
        """Set a follow-up reminder in Google Calendar"""
        if not self.service:
            print("Calendar service not initialized")
            return False
        
        if not application_date:
            application_date = datetime.now()
        elif isinstance(application_date, str):
            application_date = datetime.strptime(application_date, "%Y-%m-%d")
        
        follow_up_date = application_date + timedelta(days=follow_up_days)
        
        event = {
            'summary': f'Follow up on {job_title} application at {company_name}',
            'description': f'It\'s been {follow_up_days} days since you applied for the {job_title} position at {company_name}. Time to follow up!',
            'start': {
                'dateTime': follow_up_date.strftime("%Y-%m-%dT10:00:00"),
                'timeZone': 'America/New_York',
            },
            'end': {
                'dateTime': follow_up_date.strftime("%Y-%m-%dT10:30:00"),
                'timeZone': 'America/New_York',
            },
            'reminders': {
                'useDefault': False,
                'overrides': [
                    {'method': 'email', 'minutes': 24 * 60},
                    {'method': 'popup', 'minutes': 30},
                ],
            },
        }
        
        try:
            event = self.service.events().insert(calendarId='primary', body=event).execute()
            print(f"Follow-up reminder set for {follow_up_date.strftime('%Y-%m-%d')}")
            return True
        except Exception as e:
            print(f"Error setting reminder: {str(e)}")
            return False

# 4. Main Application Manager
class JobApplicationManager:
    def __init__(self, resume_path, cover_letter_path):
        """Initialize the job application manager"""
        self.job_retriever = JobRetriever()
        self.document_customizer = DocumentCustomizer(resume_path, cover_letter_path)
        self.calendar_manager = CalendarManager()
        self.applied_jobs = pd.DataFrame(columns=[
            'Title', 'Company', 'Location', 'URL', 'Application_Date', 
            'Resume_Path', 'CoverLetter_Path', 'FollowUp_Set'
        ])
    
    def search_jobs(self, keywords, locations):
        """Search for jobs across multiple sources and locations"""
        for location in locations:
            self.job_retriever.search_indeed(keywords, location)
            self.job_retriever.search_linkedin(keywords, location)
        
        filtered_jobs = self.job_retriever.filter_jobs()
        print(f"Found {len(filtered_jobs)} matching jobs")
        return filtered_jobs
    
    def apply_to_job(self, job_row):
        """Process a job application with customized documents and follow-up"""
        job_title = job_row['Title']
        company = job_row['Company']
        description = job_row['Description']
        
        print(f"\nProcessing application for {job_title} at {company}")
        
        # Customize resume and cover letter
        resume_path = self.document_customizer.customize_resume(job_title, company, description)
        cover_letter_path = self.document_customizer.customize_cover_letter(job_title, company, description)
        
        # Set follow-up reminder
        application_date = datetime.now()
        follow_up_set = self.calendar_manager.set_follow_up_reminder(job_title, company, application_date)
        
        # Record the application
        new_application = {
            'Title': job_title,
            'Company': company,
            'Location': job_row['Location'],
            'URL': job_row['URL'],
            'Application_Date': application_date.strftime("%Y-%m-%d"),
            'Resume_Path': resume_path,
            'CoverLetter_Path': cover_letter_path,
            'FollowUp_Set': follow_up_set
        }
        
        self.applied_jobs = pd.concat([self.applied_jobs, pd.DataFrame([new_application])], ignore_index=True)
        
        print(f"Application for {job_title} at {company} prepared successfully")
        return new_application
    
    def batch_process_jobs(self, num_jobs=5):
        """Process multiple job applications in batch"""
        filtered_jobs = self.job_retriever.filter_jobs()
        
        if len(filtered_jobs) == 0:
            print("No jobs found matching your criteria")
            return
        
        # Process the top N jobs
        jobs_to_process = filtered_jobs.head(num_jobs)
        
        for _, job in jobs_to_process.iterrows():
            self.apply_to_job(job)
        
        # Save the applied jobs record
        self.save_applied_jobs()
    
    def save_applied_jobs(self, filename='applied_jobs.csv'):
        """Save the record of applied jobs"""
        path = '/content/drive/My Drive/' + filename
        self.applied_jobs.to_csv(path, index=False)
        print(f"Saved record of {len(self.applied_jobs)} applications to {path}")
        return path

# Example usage
def run_job_application_system():
    # File paths (update these with your actual file paths)
    resume_path = "/content/drive/My Drive/Resume_Template.docx"
    cover_letter_path = "/content/drive/My Drive/Cover_Letter_Template.docx"
    
    # Initialize the job application manager
    manager = JobApplicationManager(resume_path, cover_letter_path)
    
    # Search for jobs
    job_keywords = "cybersecurity IT project management CISSP PMP"
    locations = ["New York, NY", "Remote"]
    
    jobs_found = manager.search_jobs(job_keywords, locations)
    
    # Display found jobs
    print("\nTop matching jobs:")
    for i, (_, job) in enumerate(jobs_found.head(10).iterrows()):
        print(f"{i+1}. {job['Title']} at {job['Company']} ({job['Location']}) - Match Score: {job['Keywords_Match']}")
    
    # Process applications for top 3 jobs
    manager.batch_process_jobs(num_jobs=3)
    
    print("\nJob application automation completed!")

# Run the system
run_job_application_system()
