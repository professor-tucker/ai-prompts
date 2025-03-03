 
import requests
import pandas as pd
import re
import os
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
import pickle
from google.colab import drive
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from docx import Document
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer

# Initialization
def initialize():
    drive.mount('/content/drive')
    nltk.download('punkt')
    nltk.download('stopwords')

# Job Retriever Class
class JobRetriever:
    def __init__(self):
        self.jobs_df = pd.DataFrame(columns=['Title', 'Company', 'Location', 'Description', 'URL', 'Date_Posted', 'Keywords_Match'])

    def search_indeed(self, keywords, location, pages=5):
        base_url = "https://www.indeed.com/jobs"
        all_jobs = []

        for page in range(pages):
            params = {'q': keywords, 'l': location, 'start': page * 10}
            headers = {'User-Agent': 'Mozilla/5.0'}

            try:
                response = requests.get(base_url, params=params, headers=headers)
                soup = BeautifulSoup(response.text, 'html.parser')
                job_cards = soup.find_all('div', class_='jobsearch-SerpJobCard')

                for card in job_cards:
                    job_title_elem = card.find('a', class_='jobtitle')
                    company_elem = card.find('span', class_='company')
                    location_elem = card.find('div', class_='recJobLoc')
                    description_elem = card.find('div', class_='summary')

                    if job_title_elem and company_elem:
                        job_title = job_title_elem.text.strip()
                        company = company_elem.text.strip()
                        location = location_elem['data-rc-loc'] if location_elem else "N/A"
                        description = description_elem.text.strip() if description_elem else "N/A"
                        url = "https://www.indeed.com" + job_title_elem['href']

                        keywords_list = keywords.lower().split()
                        match_score = sum(1 for keyword in keywords_list if keyword.lower() in (job_title.lower() + " " + description.lower()))

                        job_data = {'Title': job_title, 'Company': company, 'Location': location, 'Description': description, 'URL': url, 'Date_Posted': datetime.now().strftime("%Y-%m-%d"), 'Keywords_Match': match_score}
                        all_jobs.append(job_data)

            except requests.RequestException as e:
                print(f"Error scraping Indeed page {page}: {str(e)}")

        self.jobs_df = pd.concat([self.jobs_df, pd.DataFrame(all_jobs)], ignore_index=True)

    def filter_jobs(self, min_keywords_match=2):
        return self.jobs_df[self.jobs_df['Keywords_Match'] >= min_keywords_match].sort_values('Keywords_Match', ascending=False)

    def save_jobs(self, filename='job_listings.csv'):
        path = '/content/drive/My Drive/' + filename
        self.jobs_df.to_csv(path, index=False)
        print(f"Saved {len(self.jobs_df)} jobs to {path}")
        return path

# Document Customizer Class
class DocumentCustomizer:
    def __init__(self, resume_path, cover_letter_path):
        self.resume_path = resume_path
        self.cover_letter_path = cover_letter_path
        self.resume_doc = self.load_document(resume_path)
        self.cover_letter_doc = self.load_document(cover_letter_path)

    def load_document(self, path):
        try:
            return Document(path)
        except Exception as e:
            print(f"Error loading document from {path}: {str(e)}")
            return None

    def extract_job_keywords(self, job_description):
        tokens = word_tokenize(job_description.lower())
        filtered_tokens = [word for word in tokens if word.isalnum() and word not in stopwords.words('english')]
        vectorizer = TfidfVectorizer(max_features=20)
        tfidf_matrix = vectorizer.fit_transform([' '.join(filtered_tokens)])
        feature_names = vectorizer.get_feature_names_out()
        word_scores = [(word, tfidf_matrix[0, i]) for i, word in enumerate(feature_names)]
        word_scores.sort(key=lambda x: x[1], reverse=True)
        return [word for word, score in word_scores[:10]]

    def customize_resume(self, job_title, company_name, job_description):
        if not self.resume_doc:
            return None

        custom_resume = Document()
        keywords = self.extract_job_keywords(job_description)
        print(f"Keywords extracted: {keywords}")

        for para in self.resume_doc.paragraphs:
            if "[OBJECTIVE]" in para.text:
                custom_text = para.text.replace("[OBJECTIVE]", f"Experienced professional seeking the {job_title} position at {company_name}, bringing expertise in {', '.join(keywords[:3])}.")
                custom_resume.add_paragraph(custom_text, para.style)
            else:
                text = para.text
                for keyword in keywords:
                    if keyword.lower() in text.lower() and len(keyword) > 3:
                        text = re.sub(re.escape(keyword), f"**{keyword}**", text, flags=re.IGNORECASE)
                custom_resume.add_paragraph(text, para.style)

        filename = f"Custom_Resume_{company_name}_{datetime.now().strftime('%Y%m%d')}.docx"
        save_path = f"/content/drive/My Drive/{filename}"
        custom_resume.save(save_path)
        print(f"Customized resume saved to {save_path}")
        return save_path

# Initialization
initialize()

# Example usage
def run_job_application_system():
    resume_path = "/content/drive/My Drive/Resume_Template.docx"
    cover_letter_path = "/content/drive/My Drive/Cover_Letter_Template.docx"
    manager = JobApplicationManager(resume_path, cover_letter_path)
    job_keywords = "cybersecurity IT project management CISSP PMP"
    locations = ["New York, NY", "Remote"]
    jobs_found = manager.search_jobs(job_keywords, locations)
    print("\nTop matching jobs:")
    for i, (_, job) in enumerate(jobs_found.head(10).iterrows()):
        print(f"{i+1}. {job['Title']} at {job['Company']} ({job['Location']}) - Match Score: {job['Keywords_Match']}")
    manager.batch_process_jobs(num_jobs=3)
    print("\nJob application automation completed!")

run_job_application_system()
